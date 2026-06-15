from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document as DocxDocument
from pypdf import PdfReader

from app.models.schemas import ExtractedSegment
from app.utils.files import read_text_file
from app.utils.text import clean_text, join_non_empty, normalize_newlines


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".xlsx", ".md", ".txt"}


def extract_segments(path: str, original_filename: str | None = None) -> list[ExtractedSegment]:
    file_path = Path(path)
    suffix = file_path.suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Unsupported file type: {suffix}")
    if suffix == ".pdf":
        return _extract_pdf(file_path)
    if suffix == ".docx":
        return _extract_docx(file_path)
    if suffix == ".xlsx":
        return _extract_xlsx(file_path)
    return _extract_text_like(file_path, original_filename or file_path.name)


def _extract_pdf(path: Path) -> list[ExtractedSegment]:
    reader = PdfReader(str(path))
    segments: list[ExtractedSegment] = []
    for index, page in enumerate(reader.pages, start=1):
        text = clean_text(page.extract_text() or "")
        if text:
            segments.append(
                ExtractedSegment(
                    text=text,
                    source_ref=f"page {index}",
                    page_number=index,
                )
            )
    return segments


def _extract_docx(path: Path) -> list[ExtractedSegment]:
    document = DocxDocument(str(path))
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    table_lines: list[str] = []
    for table_index, table in enumerate(document.tables, start=1):
        table_lines.append(f"Table {table_index}")
        for row in table.rows:
            table_lines.append(" | ".join(cell.text.strip() for cell in row.cells))
    text = clean_text(join_non_empty([*paragraphs, *table_lines], "\n"))
    return [ExtractedSegment(text=text, source_ref="document")] if text else []


def _extract_xlsx(path: Path) -> list[ExtractedSegment]:
    workbook = pd.read_excel(path, sheet_name=None, dtype=str)
    segments: list[ExtractedSegment] = []
    for sheet_name, frame in workbook.items():
        frame = frame.fillna("")
        rows = []
        if not frame.empty:
            rows.append(" | ".join(str(col) for col in frame.columns))
            for _, row in frame.iterrows():
                rows.append(" | ".join(str(value).strip() for value in row.values if str(value).strip()))
        text = clean_text("\n".join(rows))
        if text:
            segments.append(
                ExtractedSegment(
                    text=text,
                    source_ref=f"sheet {sheet_name}",
                    sheet_name=str(sheet_name),
                )
            )
    return segments


def _extract_text_like(path: Path, original_filename: str) -> list[ExtractedSegment]:
    text = normalize_newlines(read_text_file(str(path)))
    text = clean_text(text)
    return [ExtractedSegment(text=text, source_ref=original_filename)] if text else []

